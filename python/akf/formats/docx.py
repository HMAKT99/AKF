"""AKF v1.0 — DOCX format handler.

Embeds AKF metadata into DOCX files using the OOXML Custom XML Part
mechanism. Basic embed/extract operations require NO external dependencies
(uses stdlib zipfile only). Advanced operations like auto_enrich and create
optionally use python-docx for paragraph-level introspection.

Usage:
    from akf.formats.docx import embed, extract, is_enriched, scan

    embed("report.docx", {"akf": "1.0", "claims": [...]})
    meta = extract("report.docx")
"""

from typing import List, Optional

from .base import AKFFormatHandler, ScanReport


class DOCXHandler(AKFFormatHandler):
    """Handler for Microsoft Word DOCX files."""

    FORMAT_NAME: str = "DOCX"
    EXTENSIONS: List[str] = [".docx"]
    MODE: str = "embedded"
    MECHANISM: str = "OOXML Custom XML Part"
    DEPENDENCIES: List[str] = []  # Basic ops use zipfile only

    def embed(self, filepath: str, metadata: dict) -> None:
        """Embed AKF metadata into a DOCX file.

        Uses stdlib zipfile to add a custom XML part to the OOXML archive.
        No python-docx dependency required.
        """
        from ._ooxml import embed_in_ooxml

        embed_in_ooxml(filepath, metadata)

    def extract(self, filepath: str) -> Optional[dict]:
        """Extract AKF metadata from a DOCX file.

        Returns None if no AKF metadata is embedded.
        """
        from ._ooxml import extract_from_ooxml

        return extract_from_ooxml(filepath)

    def is_enriched(self, filepath: str) -> bool:
        """Check if a DOCX file has AKF metadata embedded."""
        from ._ooxml import is_ooxml_enriched

        return is_ooxml_enriched(filepath)

    def auto_enrich(
        self,
        filepath: str,
        agent_id: str,
        default_tier: int = 3,
        classification: Optional[str] = None,
    ) -> None:
        """Auto-enrich a DOCX file with AKF metadata.

        If python-docx is available, extracts paragraph text and creates
        per-paragraph claims. Otherwise, embeds basic metadata only.
        """
        meta = self._build_auto_metadata(filepath, agent_id, default_tier, classification)

        # Try to extract paragraph content for claims if python-docx is available
        try:
            from docx import Document

            doc = Document(filepath)
            claims = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text and len(text) > 10:
                    claims.append(
                        {
                            "location": "paragraph:{}".format(i),
                            "c": text[:200],
                            "t": 0.7,
                            "src": agent_id,
                            "ai": True,
                            "tier": default_tier,
                        }
                    )
            if claims:
                meta["claims"] = claims
                meta["ai_contribution"] = 1.0
        except ImportError:
            pass  # No python-docx — embed basic metadata only

        self.embed(filepath, meta)


# ---------------------------------------------------------------------------
# Module-level convenience API
# ---------------------------------------------------------------------------

_handler = DOCXHandler()
embed = _handler.embed
extract = _handler.extract
is_enriched = _handler.is_enriched
scan = _handler.scan
auto_enrich = _handler.auto_enrich


def create(
    output: str,
    content: list,
    classification: Optional[str] = None,
    author: Optional[str] = None,
) -> None:
    """Create a new DOCX with AKF metadata from structured content.

    Requires python-docx. Install with: pip install akf[docx]

    Args:
        output: Output file path.
        content: List of content dicts. Each dict should have:
            - type: "heading" or "paragraph"
            - text: The text content
            - level: (optional) Heading level (1-6), for type="heading"
            - akf: (optional) Dict with claim metadata (t, src, ai, etc.)
        classification: Optional security classification label.
        author: Optional author/agent identifier.

    Example::

        create("report.docx", [
            {"type": "heading", "text": "Title", "level": 1},
            {"type": "paragraph", "text": "Revenue was $4.2B", "akf": {"t": 0.98}},
        ])
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "Creating DOCX requires python-docx. Install with: pip install akf[docx]"
        )

    import hashlib
    from datetime import datetime, timezone

    doc = Document()
    claims = []

    for i, item in enumerate(content):
        item_type = item.get("type", "paragraph")
        text = item.get("text", "")

        if item_type == "heading":
            doc.add_heading(text, level=item.get("level", 1))
        else:
            doc.add_paragraph(text)

        if "akf" in item:
            claim = {
                "location": "paragraph:{}".format(i),
                "c": text[:200],
            }
            claim.update(item["akf"])
            claims.append(claim)

    doc.save(output)

    # Compute hash of the saved file
    now = datetime.now(timezone.utc).isoformat()
    with open(output, "rb") as f:
        file_hash = "sha256:" + hashlib.sha256(f.read()).hexdigest()

    # Build metadata
    metadata = {
        "akf": "1.0",
        "generated_at": now,
        "overall_trust": (
            sum(c.get("t", 0.5) for c in claims) / len(claims) if claims else 0.5
        ),
        "ai_contribution": (
            sum(1 for c in claims if c.get("ai")) / len(claims) if claims else 0.0
        ),
        "claims": claims,
        "provenance": [
            {
                "actor": author or "unknown",
                "action": "created",
                "at": now,
                "hash": file_hash,
            }
        ],
        "integrity_hash": file_hash,
    }

    if classification is not None:
        metadata["classification"] = classification

    embed(output, metadata)
